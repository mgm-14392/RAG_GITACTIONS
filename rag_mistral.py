import getpass
import os
from langchain.chat_models import init_chat_model
import faiss
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.vectorstores import FAISS
from langchain import hub
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langgraph.graph import START, StateGraph
from typing_extensions import List, TypedDict
import sys
from docx import Document as DocxDocument  # For reading .docx files
from pathlib import Path

# Define State structure outside the class for accessibility
class State(TypedDict):
    question: str
    context: List[Document]
    answer: str


class RAGSystem:
    def __init__(self, openai_key=None, langsmith_key=None, mistral_key=None, docx_file_path=None):

        # Set cache location BEFORE loading any models
        cache_dir = Path("/app/model_cache")
        cache_dir.mkdir(parents=True, exist_ok=True)

        # Set cache environment variables
        os.environ["TRANSFORMERS_CACHE"] = str(cache_dir)
        os.environ["HF_HOME"] = str(cache_dir)
        os.environ["HF_DATASETS_CACHE"] = str(cache_dir)
        
        # Load the embeddings with explicit cache folder
        from langchain_huggingface import HuggingFaceEmbeddings
        
        # Set up API keys
        self.LANGSMITH_KEY = langsmith_key or os.getenv("LANGSMITH_API_KEY")
        self.MISTRAL_KEY = mistral_key or os.getenv("MISTRAL_API_KEY")

        if not self.LANGSMITH_KEY:
            raise ValueError("LANGSMITH_KEY is not set. Please set it in your environment.")
        if not self.MISTRAL_KEY:
            raise ValueError("MISTRAL_KEY is not set. Please set it in your environment.")

        os.environ["LANGSMITH_TRACING"] = "true"
        os.environ["LANGSMITH_API_KEY"] = self.LANGSMITH_KEY
        os.environ["MISTRAL_API_KEY"] = self.MISTRAL_KEY

        # Initialize components
        self.llm = init_chat_model("mistral-large-latest", model_provider="mistralai")

        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-mpnet-base-v2",
                cache_folder=str(cache_dir)  # Explicitly tell it where to look
            )
        except Exception as e:
            print(f"Error loading embeddings: {e}")
        
        self.index = faiss.IndexFlatL2(len(self.embeddings.embed_query("hello world")))
        self.vector_store = FAISS(
            embedding_function=self.embeddings,
            index=self.index,
            docstore=InMemoryDocstore(),
            index_to_docstore_id={},
        )

        # Use dynamic .docx file path if provided
        self.docx_file_path = docx_file_path or "recipes.docx"

        # Load and split documents from the .docx file
        self.docs = self.load_docx(self.docx_file_path)

        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        self.all_splits = self.text_splitter.split_documents(self.docs)

        # Index chunks
        self.vector_store.add_documents(documents=self.all_splits)

        # Load prompt template
        self.prompt = hub.pull("rlm/rag-prompt")

        # Build the application workflow
        self.graph_builder = StateGraph(State).add_sequence([self.retrieve, self.generate])
        self.graph_builder.add_edge(START, "retrieve")
        self.graph = self.graph_builder.compile()

    def load_docx(self, file_path: str) -> List[Document]:
        """
        Load and parse a .docx file into LangChain Documents.
        """
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return [Document(page_content=text)]

    def retrieve(self, state: State):
        """
        Retrieve relevant documents from vector store.
        """
        retrieved_docs = self.vector_store.similarity_search(state["question"])
        return {"context": retrieved_docs}

    def generate(self, state: State):
        """
        Generate the answer using the LLM with the retrieved context.
        """
        docs_content = "\n\n".join(doc.page_content for doc in state["context"])
        messages = self.prompt.invoke({"question": state["question"], "context": docs_content})
        response = self.llm.invoke(messages)
        return {"answer": response.content}

    def ask_question(self, question: str):
        """
        Ask a question and get an answer from the system.
        """
        response = self.graph.invoke({"question": question})
        return response["answer"]


if __name__ == "__main__":
    if len(sys.argv) > 1:
        docx_file = sys.argv[1]  # Access the argument
        print(f"Using .docx file: {docx_file}")
        # Example with dynamic .docx file path
        rag_system_default = RAGSystem(docx_file_path=docx_file)
    else:
        # Or you can still use the default file if you don't pass one
        rag_system_default = RAGSystem()
